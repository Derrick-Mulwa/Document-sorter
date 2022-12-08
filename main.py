import shutil
import docx
import os
import PyPDF2

# ASSIGN ALL REQUIRED PATHS TO RESPECTIBE VARIABLES

patha = r"C:\Users\PC\OneDrive\Documents"
pathb = r"C:\Users\PC\OneDrive\Desktop"
pathc = r"C:\Users\PC\Downloads"
docdest = r"C:\Users\PC\OneDrive\Desktop\all files\All files\All doc"
docxdest = r"C:\Users\PC\OneDrive\Desktop\all files\All files\All docx"
pdfdest = r"C:\Users\PC\OneDrive\Desktop\all files\All files\All pdf"
judgementDest = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\All Judgements"
ACC_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\ACC"
CIVIL_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\CIVIL"
ELECTION_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\ELECTION"
CRIMINAL_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\CRIMINAL"
LAND_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\LAND"
UNIDENTIFIED_Path = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\UNIDENTIFIED"
judgementDestPDF = r"C:\Users\PC\OneDrive\Desktop\all files\Judgements\All Judgements PDF"

# ASSIGN REQUIRED VARIABLES RESPECTIVE VALUES

pdf = []
doc = []
docxlist= []
uncopiedFiles = []
uncopiedJudgements = []
uncopiedPDFJudgements = []
uncopiedDocFilesList = []
uncopiedDocxFilesList = []
uncopiedPDFFilesList = []
JudgementFiles = 0
JudgementFilesPath = []
PDFJudgementFilesPath = []
readDocxFiles = 0
readPDFFiles = 0
unreadDocxFiles = 0
unidentifiedCourtList = []
unidentifiedCourt = 0
uncopiedJudgementsToCourtList = []
copiedJudgementsToCourt = 0

# IDENTIFY AND LOCATE USING PATH ALL .DOC, .DOCX AND .PDF FILES IN DOCUMENT FOLDER AND SUBFOLDERS
# PATH OF IDENTIFIED FILES IS STORED IN RESPECTIVE LISTS

for root, dirs, files in os.walk(patha):
    for f in files:
        name = os.path.join(root, f)
        ifDoc = name[-4:]
        ifDocx = name[-5:]
        ifPdf = name[-4:]

        if ifDoc == '.doc':
            doc.append(name)
        elif ifDocx == '.docx':
            docxlist.append(name)
        elif ifPdf == '.pdf':
            pdf.append(name)

# IDENTIFY AND LOCATE USING PATH ALL .DOC, .DOCX AND .PDF FILES IN DESKTOP FOLDER AND SUBFOLDERS
# PATH OF IDENTIFIED FILES IS STORED IN RESPECTIVE LISTS

for root, dirs, files in os.walk(pathb):
    for f in files:
        name = os.path.join(root, f)
        ifDoc = name[-4:]
        ifDocx = name[-5:]
        ifPdf = name[-4:]

        if ifDoc == '.doc':
            doc.append(name)
        elif ifDocx == '.docx':
            docxlist.append(name)
        elif ifPdf == '.pdf':
            pdf.append(name)

# IDENTIFY AND LOCATE USING PATH ALL .DOC, .DOCX AND .PDF FILES IN DOWNLOADS FOLDER AND SUBFOLDERS
# PATH OF IDENTIFIED FILES IS STORED IN RESPECTIVE LISTS

for root, dirs, files in os.walk(pathc):
    for f in files:
        name = os.path.join(root, f)
        ifDoc = name[-4:]
        ifDocx = name[-5:]
        ifPdf = name[-4:]

        if ifDoc == '.doc':
            doc.append(name)
        elif ifDocx == '.docx':
            docxlist.append(name)
        elif ifPdf == '.pdf':
            pdf.append(name)

# COPY ALL .DOC FILES FROM DESKTOP, DOWNLOADS, AND DOCUMENT FOLDERS TO ONE FOLDER

for docFile in doc:
    try:
        shutil.copy(docFile, docdest)
    except:
        uncopiedFiles.append(docFile)
        uncopiedDocFilesList.append(docFile)

print("COPIED ALL .DOC FILES.")

# COPY ALL .DOCX FILES FROM DESKTOP, DOWNLOADS, AND DOCUMENT FOLDERS TO ONE FOLDER

for docxFile in docxlist:
    try:
        shutil.copy(docxFile, docxdest)
    except:
        uncopiedFiles.append(docxFile)
        uncopiedDocxFilesList.append(docxFile)

print("COPIED ALL .DOCX FILES")

# COPY ALL .PDF FILES FROM DESKTOP, DOWNLOADS, AND DOCUMENT FOLDERS TO ONE FOLDER

for pdfFile in pdf:
    try:
        shutil.copy(pdfFile, pdfdest)
    except:
        uncopiedFiles.append(pdfFile)
        uncopiedPDFFilesList = []

print("COPIED ALL .PDF FILES")

# DISPLAY STATISTICS OF COPIED AND UNCOPIED FILES

print(f'Total copied files  : {len(doc)+len(docxlist)+len(pdf)} copied\n'
      f'Total uncopied files: {(len(uncopiedDocFilesList)) + (len(uncopiedDocxFilesList)) + (len(uncopiedPDFFilesList))}\n')

print("SEARCHING FOR JUDGEMENT DOCUMENTS.")

# A FUNCTION TO READ A .DOCX FILE AND IDENTIFY IF IT HAS THE WORD "JUDGEMENT" OR "RULING" IN IT.


def checkInDocx(docPath):

    doc = docx.Document(docPath)
    docxParagraphs = []

    for para in doc.paragraphs:
        docxParagraphs.append(para.text)

    if ("JUDGEMENT" in docxParagraphs) or ("JUDGMENT" in docxParagraphs) or ("RULING" in docxParagraphs):
        return True
    else:
        return False

# A FUNCTION TO READ A .PDF FILE AND IDENTIFY IF IT HAS THE WORD "JUDGEMENT" OR "RULING" IN IT.

def checkInPDF(pdfPath):

    file = open(pdfPath, "rb")
    reader = PyPDF2.PdfFileReader(file)
    page1 = reader.getPage(0)
    pdfdata = (page1.extractText()).split()


    if ("JUDGEMENT" in pdfdata) or ("RULING" in pdfdata) or ("JUDGMENT" in pdfdata):
        return True
    else:
        return False

# USE THE ABOVE FUNCTION TO CHECK ALL .DOCX FILES EARLIER COPIED AND COPY ALL DOCUMENTS WITH 'JUDGEMENT OR 'RULING'
# TO ONE FOLDER
# COPIES ALL UNCOPIED FILES' PATH TO ONE LIST


for file in docxlist:
    try:
        outcome = checkInDocx(file)
        readDocxFiles += 1

        if outcome is True:
            try:
                shutil.copy(file, judgementDest)
                JudgementFiles += 1
            except:
                uncopiedJudgements.append(file)

    except:
        unreadDocxFiles += 1

print(f"Read Docx files   : {readDocxFiles}\n"
      f"Unread Docx files : {unreadDocxFiles}")

# USE THE ABOVE FUNCTION TO CHECK ALL .PDF FILES EARLIER COPIED AND COPY ALL DOCUMENTS WITH 'JUDGEMENT OR 'RULING'
# TO ONE FOLDER
# COPIES ALL UNCOPIED FILES' PATH TO ONE LIST

for file in pdf:
    try:
        outcome = checkInPDF(file)
        readPDFFiles += 1

        if outcome is True:
            try:
                shutil.copy(file, judgementDestPDF)
                JudgementFiles += 1
            except:
                uncopiedPDFJudgements.append(file)

    except:
        unreadDocxFiles += 1

# DISPLAY STATISTICS OF THE JUDGEMENT DOCUMENTS IDENTIFIED, COPIED, AND UNCOPIED

print(f"{len(uncopiedJudgements)+JudgementFiles+(len(uncopiedPDFJudgements))} Judgement files identified. \n"
      f"Copied  : {JudgementFiles} \n"
      f"Uncopied Docx: {len(uncopiedJudgements)} \n"
      f"Uncopied PDF : {len(uncopiedPDFJudgements)}"
      f"\n{uncopiedJudgements}")

print("COPYING IDENTIFIED JUDGEMENTS TO RESPECTIVE COURTS!")

# CHECK THE COURT IN WHICH THE CASE TOOK PLACE FOR .DOCX FILES


def checkCourt(docPath):

    doc = docx.Document(docPath)
    docxParagraphs = []
    cont = ""

    for para in doc.paragraphs:
        docxParagraphs.append(para.text)

    for i in range(10):
        cont += f"{(docxParagraphs[i]).upper()} "

    requiredContent = cont.split()

    if ("ACC" in requiredContent) or ("ANTI-CORRUPTION" in requiredContent) or ("ANTI" in requiredContent) or ("EACC" in requiredContent):
        return "ACC"
    elif ("ELECTION" in requiredContent):
        return "ELECTION"
    elif ("CR" in requiredContent) or ("CR." in requiredContent) or ("CRIMINAL" in requiredContent):
        return "CRIMINAL"
    elif ("CIVIL" in requiredContent):
        return "CIVIL"
    elif ("LAND" in requiredContent):
        return "LAND"
    else:
        return "UNIDENTIFIED"

# CHECK THE COURT IN WHICH THE CASE TOOK PLACE FOR .DOCX FILES


def checkPDFCourt(pdfPath):
    file = open(pdfPath, "rb")
    reader = PyPDF2.PdfFileReader(file)
    page1 = reader.getPage(0)
    requiredContent = page1.extractText().upper().split()

    if ("ACC" in requiredContent) or ("ANTI-CORRUPTION" in requiredContent) or ("ANTI" in requiredContent) or (
            "EACC" in requiredContent):
        return "ACC"
    elif ("ELECTION" in requiredContent):
        return "ELECTION"
    elif ("CR" in requiredContent) or ("CR." in requiredContent) or ("CRIMINAL" in requiredContent):
        return "CRIMINAL"
    elif ("CIVIL" in requiredContent):
        return "CIVIL"
    elif ("LAND" in requiredContent):
        return "LAND"
    else:
        return "UNIDENTIFIED"


# COPY PATH OF ALL .DOCX JUDGEMENT DOCUMENTS TO ONE LIST

for root, dirs, files in os.walk(judgementDest):
    for f in files:
        name = os.path.join(root, f)

        JudgementFilesPath.append(name)


# COPY PATH OF ALL .PDF JUDGEMENT DOCUMENTS TO ONE LIST

for root, dirs, files in os.walk(judgementDestPDF):
    for f in files:
        name = os.path.join(root, f)

        PDFJudgementFilesPath.append(name)


# USE THE checkCourt() FUNCTION TO IDENTIFY THE COURT OF EACH .DOCX DOCUMENT AND PASTE IT IN APPROPRIATE DIRECTORY

for JudgementFile in JudgementFilesPath:
    try:
        court = checkCourt(JudgementFile)
    except:
        unidentifiedCourt += 1
        unidentifiedCourtList.append(JudgementFile)

    if court == "ACC":
        try:
            shutil.copy(JudgementFile, ACC_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "ELECTION":
        try:
            shutil.copy(JudgementFile, ELECTION_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "CRIMINAL":
        try:
            shutil.copy(JudgementFile, CRIMINAL_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "CIVIL":
        try:
            shutil.copy(JudgementFile, CIVIL_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "LAND":
        try:
            shutil.copy(JudgementFile, LAND_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "UNIDENTIFIED":
        try:
            shutil.copy(JudgementFile, UNIDENTIFIED_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

# USE THE checkPDFCourt() FUNCTION TO IDENTIFY THE COURT OF EACH .PDF DOCUMENT AND PASTE IT IN APPROPRIATE DIRECTORY

for JudgementFile in PDFJudgementFilesPath:
    try:
        court = checkPDFCourt(JudgementFile)
    except:
        unidentifiedCourt += 1
        unidentifiedCourtList.append(JudgementFile)

    if court == "ACC":
        try:
            shutil.copy(JudgementFile, ACC_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "ELECTION":
        try:
            shutil.copy(JudgementFile, ELECTION_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "CRIMINAL":
        try:
            shutil.copy(JudgementFile, CRIMINAL_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "CIVIL":
        try:
            shutil.copy(JudgementFile, CIVIL_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "LAND":
        try:
            shutil.copy(JudgementFile, LAND_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

    elif court == "UNIDENTIFIED":
        try:
            shutil.copy(JudgementFile, UNIDENTIFIED_Path)
            copiedJudgementsToCourt += 1
        except:
            uncopiedJudgementsToCourtList.append(JudgementFile)

# DISPLAY STATISTICS OF APPROPRIATE FIELDS

print(f"Unidentified Courts________________: {unidentifiedCourt}\n"
      f"Unidentified Court List____________: {unidentifiedCourtList}\n"
      f"Identified copied Judgement files  : {copiedJudgementsToCourt}\n"
      f"Identified uncopied Judgement files: {len(uncopiedJudgementsToCourtList)}\n"
      f"Identified uncopied Judgement files: {uncopiedJudgementsToCourtList}")

